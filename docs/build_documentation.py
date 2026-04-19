"""Build the bilingual Langmuir-Measurement documentation (.docx).

Run once to regenerate ``docs/LangmuirMeasure_Documentation.docx``
whenever the prose below changes.  Keeps all user-facing prose in
one self-contained script so the output is trivially reproducible.
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def _style_normal(doc):
    s = doc.styles["Normal"]
    s.font.name = "Segoe UI"
    s.font.size = Pt(11)


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Segoe UI"
    return h


def _para(doc, text, *, bold=False, italic=False, size=None,
           color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    if size is not None:
        r.font.size = size
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if color is not None:
        r.font.color.rgb = color
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    return p


def _numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    return p


def _code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    shading_elm = r._r.get_or_add_rPr()
    shd = shading_elm.makeelement(qn("w:shd"),
                                    {qn("w:val"): "clear",
                                     qn("w:color"): "auto",
                                     qn("w:fill"): "F0F2F4"})
    shading_elm.append(shd)
    return p


def _add_toc_placeholder(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = run._r.makeelement(qn("w:fldChar"),
                                   {qn("w:fldCharType"): "begin"})
    run._r.append(fldChar)
    instrText = run._r.makeelement(qn("w:instrText"),
                                     {qn("xml:space"): "preserve"})
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar2 = run._r.makeelement(qn("w:fldChar"),
                                    {qn("w:fldCharType"): "separate"})
    run._r.append(fldChar2)
    run2 = paragraph.add_run(
        "Right-click → Update Field to populate this table of contents "
        "in Microsoft Word.")
    run2.italic = True
    fldChar3 = run._r.makeelement(qn("w:fldChar"),
                                    {qn("w:fldCharType"): "end"})
    run._r.append(fldChar3)


def _table(doc, headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        c = hdr[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        cells = tbl.rows[ri].cells
        for ci, v in enumerate(row):
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(v))
            r.font.name = "Segoe UI"
            r.font.size = Pt(10)
    return tbl


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------
def build() -> Path:
    doc = Document()
    _style_normal(doc)

    # Adjust default page margins.
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---- Cover ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Langmuir Probe Measurement")
    r.bold = True
    r.font.name = "Segoe UI"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x20, 0x58, 0x8f)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(
        "Bilingual documentation — End-user manual and "
        "developer handbook")
    sr.italic = True
    sr.font.name = "Segoe UI"
    sr.font.size = Pt(14)

    _para(doc, "")
    _para(doc,
           "This document is provided in English and German.  "
           "Dieses Dokument liegt auf Englisch und Deutsch vor.",
           italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "")
    _para(doc,
           "Primary application: LPmeasurement.py on Windows lab PCs "
           "at JLU Giessen / I. Physikalisches Institut.  "
           "Produces a frozen installer LangmuirMeasure_setup.exe.",
           italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _page_break(doc)

    # ---- Table of contents ----
    _heading(doc, "Contents / Inhalt", level=1)
    _add_toc_placeholder(doc)
    _page_break(doc)

    # =======================================================================
    # Part A — short overview (bilingual)
    # =======================================================================
    _heading(doc, "A. Short overview / Kurzüberblick", level=1)

    _heading(doc, "A.1 English short overview", level=2)
    _para(doc,
           "Langmuir Probe Measurement is a Windows desktop "
           "application for Single, Double, and Triple Langmuir-probe "
           "diagnostics, plus a probe-cleaning dialog.  It drives a "
           "Keysight B2901 source-measure unit (SMU) and a Keithley "
           "2000 digital multimeter (DMM) over GPIB/USB/RS232, "
           "records the raw sweep and the derived plasma parameters "
           "to a versioned CSV, and reports the fit quality, "
           "uncertainty, and any compliance clipping directly in the "
           "operator-facing log.")
    _para(doc,
           "Quick start: install Keysight IO Libraries (or NI-VISA) "
           "on the target PC, run LangmuirMeasure_setup.exe, connect "
           "the SMU and the K2000, pick Single / Double / Triple, "
           "hit Start.  Every analysis writes a JSON sidecar next "
           "to the CSV so the exact options used are reproducible.")

    _heading(doc, "A.2 Deutscher Kurzüberblick", level=2)
    _para(doc,
           "Langmuir Probe Measurement ist eine Windows-"
           "Desktop-Anwendung für Einfach-, Doppel- und "
           "Dreifach-Langmuir-Sonden-Messungen plus einen Sonden-"
           "Reinigungsdialog.  Die Software steuert eine Keysight "
           "B2901 Source-Measure-Unit (SMU) und ein Keithley 2000 "
           "Digitalmultimeter (DMM) über GPIB/USB/RS232, speichert "
           "den Roh-Sweep und die daraus abgeleiteten Plasma-"
           "Parameter in einer versionierten CSV-Datei und zeigt "
           "Fit-Qualität, Unsicherheit sowie etwaige Compliance-"
           "Clipping-Meldungen direkt im Bedienprotokoll an.")
    _para(doc,
           "Kurzanleitung: Keysight IO Libraries (oder NI-VISA) auf "
           "dem Ziel-PC installieren, LangmuirMeasure_setup.exe "
           "ausführen, SMU und K2000 verbinden, Einfach/Doppel/"
           "Dreifach auswählen, Start klicken.  Jede Analyse "
           "schreibt eine JSON-Sidecar-Datei neben die CSV, damit "
           "die verwendeten Einstellungen später nachvollziehbar "
           "sind.")

    _page_break(doc)

    # =======================================================================
    # Part B — Full user manual (English)
    # =======================================================================
    _heading(doc, "B. Full user manual (English)", level=1)

    _heading(doc, "B.1 Purpose of the software", level=2)
    _para(doc,
           "This application turns a typical lab-bench Langmuir setup "
           "(SMU + DMM + probe) into a guided, reproducible "
           "measurement workflow.  It replaces ad-hoc SCPI scripts "
           "with a single GUI that (1) acquires a voltage sweep, "
           "(2) runs a scientifically honest analysis, (3) shows "
           "temperature / density / uncertainty / data-quality "
           "warnings, and (4) saves the data plus the exact settings "
           "so the analysis can be reproduced later.")

    _heading(doc, "B.2 Supported measurement modes", level=2)
    _table(doc,
            ["Mode", "Probe geometry", "Primary output",
             "Typical use"],
            [
                ["Single probe",
                 "One probe vs. chamber reference",
                 "T_e, V_f, V_p, n_e",
                 "Well-grounded plasmas where a wall reference is "
                 "trustworthy."],
                ["Double probe",
                 "Two identical floating tips",
                 "T_e, I_sat, n_i",
                 "RF / magnetised / floating-reference plasmas."],
                ["Triple probe",
                 "Three tips, one biased live",
                 "T_e, n_e at high time resolution",
                 "Live monitoring / fast transients."],
                ["Cleaning",
                 "Timed forced current",
                 "None (surface conditioning)",
                 "Rehabilitating a dirty probe in-situ."],
            ])

    _heading(doc, "B.3 Hardware used", level=2)
    _bullet(doc,
             "Keysight B2901A/B or B2910BL SMU (voltage source + "
             "current measurement + compliance).")
    _bullet(doc,
             "Keithley 2000 6.5-digit DMM (reference voltage readout, "
             "needed especially for Triple-probe live use).")
    _bullet(doc,
             "A GPIB-USB adapter (Keysight 82357B or NI GPIB-USB-HS) "
             "or, for the K2000 only, an RS232 link.")
    _bullet(doc,
             "The physical probe head (Single / Double / Triple tips) "
             "with matched cabling.")

    _heading(doc, "B.4 Installation prerequisites", level=2)
    _para(doc,
           "Install these on the target Windows PC **before** the "
           "application.  The frozen installer contains Python, Qt, "
           "NumPy/SciPy and the application code — but the VISA / "
           "GPIB stack cannot be redistributed and must be installed "
           "separately:")
    _numbered(doc,
               "Microsoft Visual C++ 2015–2022 x64 runtime "
               "(vc_redist.x64.exe).  Usually already present on "
               "Windows 10 / 11; the installer can chain it if a "
               "copy is staged next to the .iss file.")
    _numbered(doc,
               "A system VISA library.  Pick either: Keysight IO "
               "Libraries Suite (recommended for Keysight SMUs) "
               "or NI-VISA (best for NI GPIB-USB adapters).  "
               "Either provides visa32.dll / visa64.dll plus the "
               "GPIB kernel driver for the chosen adapter.")
    _numbered(doc,
               "A USB-to-RS232 driver, only if the K2000 is wired "
               "through a USB converter (FTDI, Prolific, etc.).")
    _para(doc,
           "Full checklist: see docs/INSTALL_prereqs.md bundled in "
           "the repository.  The installer also shows a warning "
           "dialog if no VISA library is detected.")

    _heading(doc, "B.5 Installing the application", level=2)
    _numbered(doc,
               "Double-click LangmuirMeasure_setup.exe.")
    _numbered(doc,
               "If prompted that no VISA library was detected, "
               "finish the VISA installer first, then re-run the "
               "setup.")
    _numbered(doc,
               "Follow the wizard; accept the default install "
               "folder unless you have a specific reason to change "
               "it.  A Start-menu shortcut and (optionally) a "
               "desktop icon are created.")
    _numbered(doc,
               "Launch Langmuir Probe Measurement from the Start "
               "menu.  The main window opens maximised.")

    _heading(doc, "B.6 First startup", level=2)
    _para(doc,
           "The first launch creates a writable user-data folder "
           "under %APPDATA%\\JLU-IPI\\DLP\\ that stores the VISA "
           "scan cache, the persistent analysis history, and the "
           "per-method measurement sub-folders.  Nothing is written "
           "into %ProgramFiles%, so the application runs under a "
           "standard (non-administrator) Windows user account.")
    _para(doc,
           "You will see three columns: left controls, centre plot + "
           "method selector, right K2000 panel + log.  Column widths "
           "can be dragged with the splitter handles.")

    _heading(doc, "B.7 Connecting the instruments", level=2)
    _para(doc, "SMU (Keysight B2901 / B2910BL):")
    _numbered(doc,
               "Pick the VISA resource in the combo box (e.g. "
               "GPIB0::23::INSTR).  If empty, press Scan; the app "
               "enumerates all reachable instruments and remembers "
               "the last-successful selection.")
    _numbered(doc,
               "Click Connect.  A green LED + the instrument's "
               "*IDN? string indicates success.  Any failure is "
               "surfaced with a classified reason: install VISA, "
               "check GPIB address, check cables, or a timeout "
               "remediation hint.")
    _para(doc, "Keithley 2000 multimeter:")
    _numbered(doc,
               "In the right-column K2000 panel choose Transport "
               "= GPIB (default) or RS232, enter the address / COM "
               "port, and click Connect.")
    _numbered(doc,
               "Optionally open K2000 Options… to configure "
               "autorange, fixed voltage range, and NPLC "
               "integration time.  Defaults are autorange ON, "
               "NPLC 1.0.  For Triple-probe live use, autorange "
               "OFF + NPLC 0.1 is a typical choice.")

    _heading(doc, "B.8 Walking through the GUI", level=2)
    _para(doc,
           "The left control column groups the sweep parameters: "
           "Voltage start / stop / step, Settle time, Averages, "
           "Bidirectional, Compliance (current limit), Save CSV, "
           "Auto-analyze.  Below the plot, the method band lets you "
           "switch between Single, Double, Triple, and Cleaning; the "
           "button Probe Params… lets you set probe geometry.  The "
           "right column shows the K2000 panel and a compact log.")
    _para(doc, "Key buttons:")
    _table(doc,
            ["Button", "What it does"],
            [
                ["Scan",
                 "Enumerate reachable VISA instruments."],
                ["Connect / Disconnect",
                 "Open or close the SMU link."],
                ["Start / Stop",
                 "Begin or abort a voltage sweep."],
                ["Analyze",
                 "Run the Single / Double / Triple pipeline on the "
                 "currently loaded sweep."],
                ["Probe Params…",
                 "Geometry (cylinder / sphere / planar), electrode "
                 "length, radius, area."],
                ["Experiment…",
                 "Gas species and flow rates; drives the ion-mass "
                 "used in the density calculation."],
                ["Fit Model…",
                 "Combined dialog for fit-model choice, compliance "
                 "handling, hysteresis threshold, bootstrap CI, and "
                 "a Help button that opens the Double-probe "
                 "documentation."],
                ["Instrument…",
                 "SMU settings: NPLC, autorange, output protection, "
                 "remote sense, ..."],
                ["K2000 Options…",
                 "Autorange, fixed V range, NPLC for the Keithley "
                 "2000."],
                ["Plot…",
                 "Axis range / grid / legend / reset view."],
            ])

    _heading(doc, "B.9 How Single / Double / Triple analysis works",
              level=2)
    _para(doc,
           "Single-probe analysis takes the I–V curve, finds the "
           "floating potential V_f as the zero crossing, estimates "
           "the ion-saturation branch, refines the electron "
           "temperature T_e by a semilog fit on the retarding "
           "region, and finally derives V_p (plasma potential) and "
           "n_e (electron density from the Bohm flux).  A robust "
           "Huber-loss fit is used by default.  Optional: a non-"
           "parametric bootstrap gives a numeric 95 % CI for T_e.")
    _para(doc,
           "Double-probe analysis fits a tanh-family model to the "
           "sweep: I(V) = I_sat · tanh(V/W) [+ g·V] [· (1 + a·tanh)]. "
           "T_e = W / 2.  Three model variants are available "
           "(simple, tanh+slope, tanh+slope+asymmetry).  A "
           "conservative compliance guard removes clipped points or "
           "marks them in the result; the always-on covariance-based "
           "95 % CI for T_e, I_sat and n_i (fit-only) can be "
           "replaced by a residual-bootstrap CI.  For legacy CSVs "
           "without a compliance column, a plateau heuristic may "
           "detect suspected clipping.")
    _para(doc,
           "Triple-probe live measurement streams samples in a "
           "worker thread and computes T_e and n_e on each tick "
           "from the instantaneous ratio of the two potential "
           "differences.  Values appear in the Triple window with "
           "a rolling plot and are saved to a method-tagged CSV.")

    _heading(doc, "B.10 Interpreting the output", level=2)
    _table(doc,
            ["Field", "Meaning"],
            [
                ["T_e ± σ",
                 "Electron temperature (eV) and its 1-σ "
                 "uncertainty from the fit covariance."],
                ["T_e 95 % CI",
                 "Covariance-based interval ±1.96 σ, or — when "
                 "bootstrap is enabled — the non-parametric "
                 "percentile interval."],
                ["I_sat",
                 "Ion saturation current, always shown with its "
                 "own covariance-based 95 % CI."],
                ["n_i (fit-only)",
                 "Bohm-flux density derived from I_sat + T_e + "
                 "probe area + ion mass.  The CI is labelled "
                 "fit-only because probe-area and ion-mass "
                 "uncertainty are treated as exact inputs."],
                ["Fit [grade]",
                 "excellent / good / fair / poor, derived from "
                 "R² and NRMSE thresholds."],
                ["Status",
                 "OK, POOR, WARNING, or a failure (non_converged, "
                 "bounds_error, insufficient_data, bad_input, "
                 "numerical_error).  Failure statuses mean the "
                 "numbers shown are NOT trustworthy."],
                ["Compliance row",
                 "Appears whenever clipping was detected: N/M "
                 "flagged points, percentage, and the action "
                 "taken (excluded vs retained).  Labels legacy-"
                 "heuristic findings as 'suspected clipping'."],
            ])

    _heading(doc, "B.11 Warnings and common failure messages",
              level=2)
    _bullet(doc,
             "Fit failed — non_converged: the optimiser hit maxfev; "
             "check the sweep range and the sign of the current.")
    _bullet(doc,
             "Fit failed — insufficient_data: too few points or no "
             "voltage variance; increase the sweep range or the "
             "number of steps.")
    _bullet(doc,
             "Fit warning — clipping retained: the chosen "
             "Compliance = Include all mode kept clipped points in "
             "the fit; either re-acquire with a higher compliance "
             "limit or switch to Exclude clipped.")
    _bullet(doc,
             "K2000 connect failed — no_visa: install Keysight IO "
             "Libraries or NI-VISA and try again.")
    _bullet(doc,
             "K2000 connect failed — no_device: check the GPIB / "
             "RS232 address and that the instrument is powered and "
             "visible in Keysight Connection Expert / NI MAX.")
    _bullet(doc,
             "Analysis sidecar write failed: non-fatal.  The "
             "numbers on screen are still good; only the per-run "
             "JSON sidecar next to the CSV could not be written.")

    _heading(doc, "B.12 Saving, loading, and sidecar files",
              level=2)
    _para(doc,
           "Every sweep can be saved to a CSV.  Files are organised "
           "per method: <base>/single/, <base>/double/, "
           "<base>/triple/.  The filename pattern is "
           "LP_<ISO timestamp>_<method>.csv.  The CSV header starts "
           "with a short banner plus a schema marker "
           "(Schema: lp-measurement-csv v1); legacy files without "
           "this marker still load.")
    _para(doc,
           "When you click Analyze, a JSON sidecar file is written "
           "next to the CSV: <stem>.options.json.  It captures the "
           "analysis options (compliance mode, bootstrap, "
           "hysteresis threshold, ...), the fit model, and a "
           "summary of the numeric result (T_e, CIs, status).  This "
           "makes later re-analysis auditable.")

    _heading(doc, "B.13 Practical troubleshooting", level=2)
    _bullet(doc,
             "Nothing happens on Start: make sure SMU is connected "
             "(green LED) AND the method is not Cleaning.")
    _bullet(doc,
             "K2000 reads but always returns 0.6 V: the Sim "
             "checkbox is on — uncheck it for real hardware.")
    _bullet(doc,
             "GPIB was working yesterday, now VI_ERROR_RSRC_NFOUND: "
             "unplug and re-plug the USB-GPIB adapter, run Scan "
             "again; the GPIB primary address on the instrument may "
             "have changed.")
    _bullet(doc,
             "Analyze shows Status = POOR even on clean-looking "
             "data: check the warning reason in the result block — "
             "it often points at clipping or a model mismatch "
             "(switch model variant).")
    _bullet(doc,
             "Installer refused to start with 'no VISA library' "
             "warning: you can still continue; RS232 + TCP paths "
             "may still work.  GPIB will not.")

    _heading(doc, "B.14 Recommended operator workflow", level=2)
    _numbered(doc,
               "Power up SMU + K2000 + probe chamber.  Wait 1 min "
               "for the SMU to warm up.")
    _numbered(doc,
               "Open the application, Scan, Connect both "
               "instruments.  Verify the IDN strings.")
    _numbered(doc,
               "Set probe geometry (Probe Params…) and gas "
               "(Experiment…) if you want a density number.")
    _numbered(doc,
               "Pick the method (Single / Double / Triple / "
               "Cleaning).")
    _numbered(doc,
               "Set sweep parameters (V_start, V_stop, V_step, "
               "Compliance, Averages, Settle).")
    _numbered(doc,
               "Press Start.  Wait for the sweep to finish; the "
               "Save CSV checkbox saves it automatically.")
    _numbered(doc,
               "Press Analyze.  Inspect the compact result block "
               "for Status, CIs and any Compliance row.  If the "
               "result looks wrong, try a different fit model or "
               "re-acquire with a higher compliance limit.")

    _page_break(doc)

    # =======================================================================
    # Part C — Full user manual (German)
    # =======================================================================
    _heading(doc, "C. Vollständiges Benutzerhandbuch (Deutsch)",
              level=1)

    _heading(doc, "C.1 Zweck der Software", level=2)
    _para(doc,
           "Diese Anwendung macht aus einem typischen Langmuir-"
           "Aufbau (SMU + DMM + Sonde) einen geführten, "
           "reproduzierbaren Messablauf.  Sie ersetzt ad-hoc "
           "SCPI-Skripte durch eine einzige grafische Oberfläche, "
           "die (1) einen Spannungs-Sweep aufnimmt, (2) eine "
           "wissenschaftlich ehrliche Auswertung durchführt, "
           "(3) Temperatur, Dichte, Unsicherheit und Daten-"
           "qualitäts-Warnungen anzeigt, und (4) die Daten samt "
           "verwendeter Einstellungen speichert, damit die Analyse "
           "später nachvollziehbar bleibt.")

    _heading(doc, "C.2 Unterstützte Messmodi", level=2)
    _table(doc,
            ["Modus", "Sondengeometrie", "Hauptausgabe",
             "Typische Anwendung"],
            [
                ["Einfachsonde",
                 "Eine Sonde gegen Kammerreferenz",
                 "T_e, V_f, V_p, n_e",
                 "Gut geerdete Plasmen mit vertrauenswürdiger "
                 "Wandreferenz."],
                ["Doppelsonde",
                 "Zwei baugleiche, schwebende Spitzen",
                 "T_e, I_sat, n_i",
                 "RF- / magnetisierte / schwebend referenzierte "
                 "Plasmen."],
                ["Dreifachsonde",
                 "Drei Spitzen, eine davon aktiv vorgespannt",
                 "T_e, n_e in hoher Zeitauflösung",
                 "Live-Monitoring und schnelle Transienten."],
                ["Reinigung",
                 "Zeitgesteuerter Prüfstrom",
                 "Keine (Oberflächen-Konditionierung)",
                 "Reinigung einer verschmutzten Sonde im Betrieb."],
            ])

    _heading(doc, "C.3 Verwendete Hardware", level=2)
    _bullet(doc,
             "Keysight B2901A/B oder B2910BL SMU (Spannungsquelle + "
             "Strommessung + Compliance).")
    _bullet(doc,
             "Keithley 2000 6½-stelliges DMM (Referenz-Spannungs-"
             "messung, insbesondere für Dreifachsonden-Live-Betrieb).")
    _bullet(doc,
             "Ein GPIB-USB-Adapter (Keysight 82357B oder NI "
             "GPIB-USB-HS); für den K2000 alternativ RS232.")
    _bullet(doc,
             "Der physische Sondenkopf (Einfach / Doppel / "
             "Dreifach) mit passenden Kabeln.")

    _heading(doc, "C.4 Installationsvoraussetzungen", level=2)
    _para(doc,
           "Diese Komponenten müssen auf dem Ziel-Windows-PC "
           "**vor** der Anwendung installiert sein.  Der gefrorene "
           "Installer enthält Python, Qt, NumPy/SciPy und die "
           "Applikation — der VISA/GPIB-Stack darf aber lizenz-"
           "bedingt nicht mitgeliefert werden und wird separat "
           "installiert:")
    _numbered(doc,
               "Microsoft Visual C++ 2015–2022 x64 Runtime "
               "(vc_redist.x64.exe).  Unter Windows 10 / 11 "
               "meistens bereits vorhanden; der Installer kann sie "
               "optional silent mitinstallieren, wenn die Datei "
               "neben der .iss-Datei liegt.")
    _numbered(doc,
               "Eine System-VISA-Bibliothek.  Wahlweise: "
               "Keysight IO Libraries Suite (empfohlen für "
               "Keysight-SMUs) oder NI-VISA (am besten bei "
               "NI-GPIB-USB-Adaptern).  Beide liefern visa32.dll / "
               "visa64.dll plus den GPIB-Kernel-Treiber für den "
               "gewählten Adapter.")
    _numbered(doc,
               "Ein USB-zu-RS232-Treiber, nur falls der K2000 "
               "über einen USB-Seriell-Adapter angebunden ist "
               "(FTDI, Prolific, o. ä.).")
    _para(doc,
           "Vollständige Checkliste: docs/INSTALL_prereqs.md im "
           "Repository.  Der Installer warnt beim Start, wenn keine "
           "VISA-Bibliothek gefunden wird.")

    _heading(doc, "C.5 Installation der Anwendung", level=2)
    _numbered(doc,
               "LangmuirMeasure_setup.exe per Doppelklick starten.")
    _numbered(doc,
               "Erscheint die Warnung, dass keine VISA-Bibliothek "
               "gefunden wurde: zuerst die VISA-Installation "
               "abschließen und das Setup erneut aufrufen.")
    _numbered(doc,
               "Dem Assistenten folgen; den Standard-"
               "Installationspfad übernehmen, sofern kein "
               "konkreter Grund dagegen spricht.  Ein Startmenü-"
               "Eintrag und optional ein Desktop-Symbol werden "
               "angelegt.")
    _numbered(doc,
               "Langmuir Probe Measurement aus dem Startmenü "
               "starten.  Das Hauptfenster öffnet maximiert.")

    _heading(doc, "C.6 Erster Start", level=2)
    _para(doc,
           "Beim ersten Start legt die Anwendung ein beschreibbares "
           "Benutzerdaten-Verzeichnis unter %APPDATA%\\JLU-IPI\\DLP\\ "
           "an, in dem der VISA-Scan-Cache, die Analyse-Historie "
           "und die messmethoden-spezifischen Unterordner liegen.  "
           "Im %ProgramFiles% wird nichts geschrieben, sodass die "
           "Software auch ohne Administrator-Rechte läuft.")
    _para(doc,
           "Drei Spalten sind sichtbar: links die Bedienelemente, "
           "in der Mitte der Plot samt Methodenwahl, rechts das "
           "K2000-Panel und das Log.  Die Spaltenbreiten können "
           "mit den Splitter-Griffen angepasst werden.")

    _heading(doc, "C.7 Geräte verbinden", level=2)
    _para(doc, "SMU (Keysight B2901 / B2910BL):")
    _numbered(doc,
               "VISA-Ressource in der Combobox auswählen "
               "(z. B. GPIB0::23::INSTR).  Ist die Liste leer: "
               "Scan klicken.  Die Anwendung enumeriert alle "
               "erreichbaren Instrumente und merkt sich die "
               "zuletzt erfolgreich verwendete Auswahl.")
    _numbered(doc,
               "Connect klicken.  Grüne LED + *IDN?-Antwort "
               "zeigen Erfolg.  Fehler werden klassifiziert "
               "gemeldet: VISA fehlt, falsche Adresse, Kabel, "
               "Timeout — jeweils mit konkretem Hinweis.")
    _para(doc, "Keithley 2000 Multimeter:")
    _numbered(doc,
               "Im K2000-Panel rechts Transport = GPIB (Standard) "
               "oder RS232 wählen, Adresse / COM-Port eintragen, "
               "Connect klicken.")
    _numbered(doc,
               "Optional K2000 Options… öffnen, um Autorange, "
               "fixen Spannungsbereich und NPLC einzustellen.  "
               "Standards: Autorange EIN, NPLC 1.0.  Für den "
               "Dreifachsonden-Live-Betrieb ist Autorange AUS + "
               "NPLC 0.1 ein typischer Kompromiss zwischen "
               "Geschwindigkeit und Rauschen.")

    _heading(doc, "C.8 Die Bedienoberfläche im Detail", level=2)
    _para(doc,
           "Die linke Bedienspalte gruppiert die Sweep-Parameter: "
           "Spannung Start / Stop / Schritt, Settle-Zeit, "
           "Mittelungen, bidirektional, Compliance (Strombegrenzung), "
           "CSV speichern, Auto-analyze.  Unterhalb des Plots lässt "
           "die Methodenleiste zwischen Einfach, Doppel, Dreifach "
           "und Reinigung wählen; die Schaltfläche Probe Params… "
           "öffnet die Sondengeometrie.  Rechts: K2000-Panel und "
           "kompaktes Log.")
    _para(doc, "Wichtige Schaltflächen:")
    _table(doc,
            ["Schaltfläche", "Funktion"],
            [
                ["Scan",
                 "Erreichbare VISA-Instrumente auflisten."],
                ["Connect / Disconnect",
                 "Verbindung zur SMU auf- bzw. abbauen."],
                ["Start / Stop",
                 "Spannungs-Sweep starten bzw. abbrechen."],
                ["Analyze",
                 "Pipeline für Einfach / Doppel / Dreifach auf "
                 "den aktuellen Sweep anwenden."],
                ["Probe Params…",
                 "Geometrie (Zylinder / Kugel / Planar), "
                 "Elektrodenlänge, Radius, Fläche."],
                ["Experiment…",
                 "Gasart und Fluss; liefert die Ionenmasse für "
                 "die Dichtebestimmung."],
                ["Fit Model…",
                 "Kombinierter Dialog für Fit-Modell, Compliance-"
                 "Verhalten, Hysterese-Schwelle, Bootstrap-CI und "
                 "eine Help-Schaltfläche zur Doppelsonden-"
                 "Dokumentation."],
                ["Instrument…",
                 "SMU-Einstellungen: NPLC, Autorange, Ausgangs-"
                 "schutz, Remote-Sense, ..."],
                ["K2000 Options…",
                 "Autorange, fixer Spannungsbereich, NPLC für "
                 "das Keithley 2000."],
                ["Plot…",
                 "Achsenbereich / Gitter / Legende / Ansicht "
                 "zurücksetzen."],
            ])

    _heading(doc, "C.9 Funktionsweise der Analysen", level=2)
    _para(doc,
           "Einfachsonden-Analyse: aus der I–V-Kennlinie wird "
           "V_f (Floating-Potential) als Nulldurchgang bestimmt, "
           "die Ionen-Sättigungsbranche geschätzt, T_e per semilog-"
           "Fit in der Retardierungszone verfeinert und schließlich "
           "V_p (Plasmapotential) sowie n_e (Elektronen-dichte "
           "über den Bohm-Fluss) abgeleitet.  Standardmäßig wird "
           "ein robuster Huber-Fit verwendet.  Optional liefert "
           "ein nicht-parametrischer Bootstrap ein numerisches "
           "95 %-Konfidenzintervall für T_e.")
    _para(doc,
           "Doppelsonden-Analyse: ein Tanh-Modell wird an den "
           "Sweep angepasst: I(V) = I_sat · tanh(V/W) [+ g·V] "
           "[· (1 + a·tanh)].  T_e = W / 2.  Drei Varianten sind "
           "verfügbar (einfach, tanh+Steigung, tanh+Steigung+"
           "Asymmetrie).  Ein konservativer Compliance-Schutz "
           "entfernt geclippte Punkte oder kennzeichnet sie; die "
           "stets vorhandene kovarianzbasierte 95 %-CI für T_e, "
           "I_sat und n_i (nur Fit-Unsicherheit) kann durch einen "
           "Residuen-Bootstrap ersetzt werden.  Bei Legacy-CSVs "
           "ohne Compliance-Spalte erkennt eine Plateau-Heuristik "
           "vermutetes Clipping.")
    _para(doc,
           "Dreifachsonden-Live-Messung: ein Worker-Thread nimmt "
           "Proben und berechnet T_e und n_e tickweise aus dem "
           "Verhältnis der beiden Potentialdifferenzen.  Die "
           "Werte erscheinen im Triple-Fenster mit rollendem Plot "
           "und werden in einer methoden-getaggten CSV gespeichert.")

    _heading(doc, "C.10 Ergebnisse interpretieren", level=2)
    _table(doc,
            ["Feld", "Bedeutung"],
            [
                ["T_e ± σ",
                 "Elektronentemperatur (eV) mit 1-σ-Unsicherheit "
                 "aus der Fit-Kovarianz."],
                ["T_e 95 % CI",
                 "Kovarianzbasiertes Intervall ±1,96 σ, oder — "
                 "bei aktiviertem Bootstrap — das nicht-"
                 "parametrische Perzentil-Intervall."],
                ["I_sat",
                 "Ionen-Sättigungsstrom, stets mit eigenem "
                 "kovarianzbasiertem 95 %-CI."],
                ["n_i (fit-only)",
                 "Bohm-Fluss-Dichte aus I_sat + T_e + Sondenfläche "
                 "+ Ionenmasse.  Die CI ist als „fit-only“ "
                 "gekennzeichnet, weil Sondenflächen- und Ionen-"
                 "massen-Unsicherheiten als exakt behandelt werden."],
                ["Fit [grade]",
                 "exzellent / gut / mittel / schlecht, abgeleitet "
                 "aus R²- und NRMSE-Schwellen."],
                ["Status",
                 "OK, POOR, WARNING oder ein Fehler "
                 "(non_converged, bounds_error, insufficient_data, "
                 "bad_input, numerical_error).  Fehler-Status "
                 "bedeutet: die Zahlen sind nicht vertrauenswürdig."],
                ["Compliance-Zeile",
                 "Erscheint bei erkanntem Clipping: N/M markierte "
                 "Punkte, Prozentsatz, durchgeführte Aktion "
                 "(ausgeschlossen vs. belassen).  Legacy-Heuristik-"
                 "Treffer werden als „suspected clipping“ "
                 "gekennzeichnet."],
            ])

    _heading(doc, "C.11 Warnungen und typische Fehlermeldungen",
              level=2)
    _bullet(doc,
             "Fit failed — non_converged: der Optimierer hat "
             "maxfev erreicht; Sweep-Bereich und Vorzeichen des "
             "Stroms prüfen.")
    _bullet(doc,
             "Fit failed — insufficient_data: zu wenige Punkte "
             "oder keine Spannungsänderung; Sweep verlängern oder "
             "mehr Schritte.")
    _bullet(doc,
             "Fit warning — clipping retained: Compliance = "
             "„Include all“ hat geclippte Punkte im Fit "
             "belassen; mit höherem Compliance-Limit erneut "
             "messen oder auf „Exclude clipped“ umstellen.")
    _bullet(doc,
             "K2000 connect failed — no_visa: Keysight IO "
             "Libraries oder NI-VISA installieren und erneut "
             "versuchen.")
    _bullet(doc,
             "K2000 connect failed — no_device: GPIB-/RS232-"
             "Adresse prüfen, Gerät eingeschaltet, in Keysight "
             "Connection Expert / NI MAX sichtbar?")
    _bullet(doc,
             "Analysis sidecar write failed: nicht fatal.  Die "
             "angezeigten Zahlen sind gültig; nur die JSON-"
             "Sidecar-Datei neben der CSV konnte nicht "
             "geschrieben werden.")

    _heading(doc, "C.12 Speichern, Laden, Sidecar-Dateien",
              level=2)
    _para(doc,
           "Jeder Sweep kann als CSV gespeichert werden.  Dateien "
           "werden methoden-getrennt abgelegt: <base>/single/, "
           "<base>/double/, <base>/triple/.  Das Dateinamenschema "
           "lautet LP_<ISO-Zeitstempel>_<methode>.csv.  Der "
           "CSV-Kopf beginnt mit einem kurzen Banner plus einer "
           "Schema-Kennung (Schema: lp-measurement-csv v1); "
           "Legacy-Dateien ohne diese Kennung lassen sich "
           "weiterhin laden.")
    _para(doc,
           "Nach einem Klick auf Analyze wird neben der CSV eine "
           "JSON-Sidecar-Datei geschrieben: <stem>.options.json. "
           "Sie enthält die Analyse-Optionen (Compliance-Modus, "
           "Bootstrap, Hysterese-Schwelle, ...), das Fit-Modell "
           "und eine Zusammenfassung des Ergebnisses (T_e, CIs, "
           "Status).  So bleibt eine spätere Re-Analyse "
           "nachvollziehbar.")

    _heading(doc, "C.13 Praktische Fehlersuche", level=2)
    _bullet(doc,
             "Start reagiert nicht: SMU verbunden (grüne LED) "
             "und Methode ist nicht „Reinigung“?")
    _bullet(doc,
             "K2000 liefert immer 0,6 V: das Sim-Häkchen ist "
             "aktiv; für echte Hardware deaktivieren.")
    _bullet(doc,
             "GPIB ging gestern, heute VI_ERROR_RSRC_NFOUND: "
             "USB-GPIB-Adapter aus- und einstecken, Scan erneut; "
             "die primäre GPIB-Adresse am Gerät kann sich "
             "geändert haben.")
    _bullet(doc,
             "Analyze zeigt Status = POOR trotz sauberer Kurve: "
             "Warnungs-Grund im Ergebnisblock lesen — meist "
             "Clipping oder ein Modell-Missfit (andere Variante "
             "probieren).")
    _bullet(doc,
             "Installer startet mit Warnung „keine VISA-"
             "Bibliothek“: Installation kann fortgesetzt werden; "
             "RS232- und TCP-Pfade funktionieren evtl. trotzdem. "
             "GPIB nicht.")

    _heading(doc, "C.14 Empfohlener Arbeitsablauf", level=2)
    _numbered(doc,
               "SMU + K2000 + Sondenkammer einschalten.  1 Minute "
               "Aufwärmzeit für die SMU.")
    _numbered(doc,
               "Anwendung öffnen, Scan, beide Geräte verbinden.  "
               "IDN-Strings prüfen.")
    _numbered(doc,
               "Sondengeometrie (Probe Params…) und Gas "
               "(Experiment…) setzen, falls eine Dichte gewünscht "
               "ist.")
    _numbered(doc,
               "Methode wählen (Einfach / Doppel / Dreifach / "
               "Reinigung).")
    _numbered(doc,
               "Sweep-Parameter setzen (V_start, V_stop, V_step, "
               "Compliance, Mittelungen, Settle).")
    _numbered(doc,
               "Start drücken.  Auf Fertigstellung warten; "
               "„Save CSV“ speichert automatisch.")
    _numbered(doc,
               "Analyze drücken.  Kompaktes Ergebnis-Panel "
               "prüfen (Status, CIs, Compliance-Zeile).  Wirkt "
               "das Ergebnis falsch: anderes Fit-Modell oder "
               "höheres Compliance-Limit versuchen.")

    _page_break(doc)

    # =======================================================================
    # Part D — developer / GitHub documentation (bilingual blocks)
    # =======================================================================
    _heading(doc, "D. Developer / GitHub documentation",
              level=1)

    _heading(doc, "D.1 Project purpose (EN)", level=2)
    _para(doc,
           "A focused desktop application for Langmuir-probe "
           "diagnostics on Windows lab PCs, written in Python / "
           "PySide6 with strict separation between GUI code and "
           "pure-Python analysis code.  Packaged via PyInstaller "
           "and Inno Setup.  Designed for small-team operation "
           "with rigorous test coverage and reproducible analysis "
           "artefacts.")
    _heading(doc, "D.1 Projektzweck (DE)", level=2)
    _para(doc,
           "Eine fokussierte Desktop-Anwendung für Langmuir-"
           "Sonden-Diagnostik auf Windows-Laborrechnern, "
           "geschrieben in Python / PySide6 mit strikter Trennung "
           "zwischen GUI- und reinem Analyse-Code.  Paketierung "
           "über PyInstaller und Inno Setup.  Entwickelt für "
           "kleine Teams mit solider Testabdeckung und "
           "reproduzierbaren Analyse-Artefakten.")

    _heading(doc, "D.2 Architecture overview / "
              "Architekturüberblick", level=2)
    _para(doc,
           "EN — Layers: (1) instrument drivers "
           "(keysight_b2901.py, keithley_2000.py) expose a small, "
           "testable surface; (2) analysis modules "
           "(dlp_single_analysis, dlp_double_analysis, "
           "dlp_triple_analysis, dlp_fit_models) are pure "
           "numpy/scipy; (3) GUI modules (LPmeasurement.py and its "
           "dialogs) import Qt only at runtime; (4) persistence / "
           "paths / CSV schema / sidecar / VISA cache helpers sit "
           "underneath everything.  The entry point is "
           "LPmeasurement.py, which inherits from the V2 window "
           "for legacy widget re-use.")
    _para(doc,
           "DE — Schichten: (1) Gerätetreiber "
           "(keysight_b2901.py, keithley_2000.py) mit kleiner, "
           "testbarer Schnittstelle; (2) Auswertungsmodule "
           "(dlp_single_analysis, dlp_double_analysis, "
           "dlp_triple_analysis, dlp_fit_models) rein in "
           "numpy/scipy; (3) GUI-Module (LPmeasurement.py und "
           "zugehörige Dialoge) importieren Qt nur zur Laufzeit; "
           "(4) Persistenz / Pfade / CSV-Schema / Sidecar / "
           "VISA-Cache bilden das Fundament.  Einstiegspunkt ist "
           "LPmeasurement.py, das von der V2-Fensterklasse erbt, "
           "um das Layout wiederzuverwenden.")

    _heading(doc, "D.3 Key modules", level=2)
    _table(doc,
            ["Module", "Responsibility"],
            [
                ["LPmeasurement.py",
                 "Main window + method dispatcher + override of "
                 "V2 analyze / save paths."],
                ["DoubleLangmuir_measure_v2.py",
                 "Shared widget layout, worker thread, V2 analyze "
                 "entry point (also called by LP's override)."],
                ["dlp_single_analysis.py",
                 "Single-probe pipeline: V_f, T_e, V_p, n_e + "
                 "bootstrap CI."],
                ["dlp_double_analysis.py",
                 "Double-probe orchestrator: saturation fit → "
                 "tanh model → density → clipping guard."],
                ["dlp_triple_analysis.py / dlp_triple_worker.py",
                 "Triple-probe live sampling + ratio math."],
                ["dlp_fit_models.py",
                 "Tanh model registry, fit_dlp_model (with status), "
                 "bootstrap_te_ci_double."],
                ["dlp_double_options.py / dlp_single_options.py",
                 "Dataclasses + dialogs for per-method analysis "
                 "options."],
                ["dlp_double_help.py / dlp_single_help.py",
                 "Rich-text help dialogs shared between Single and "
                 "Double."],
                ["dlp_k2000_options.py",
                 "K2000 autorange / range / NPLC options."],
                ["keysight_b2901.py / keithley_2000.py / "
                 "fake_*.py",
                 "Real + simulated instrument drivers, both "
                 "returning ClassifiedVisaError."],
                ["visa_errors.py",
                 "VISA error classification + remediation hints."],
                ["dlp_csv_schema.py",
                 "CSV schema v1 banner (Schema: lp-measurement-csv "
                 "v1)."],
                ["analysis_options_sidecar.py",
                 "Per-analysis JSON sidecar reader/writer."],
                ["clipping_heuristic.py",
                 "Plateau heuristic for legacy datasets without "
                 "compliance column."],
                ["paths.py",
                 "Frozen-safe user-data paths under "
                 "%APPDATA%/JLU-IPI/DLP."],
                ["LangmuirMeasure.spec + build.bat + "
                 "LangmuirMeasure_setup.iss",
                 "PyInstaller spec, build driver, Inno Setup "
                 "script (with optional VC++ chain)."],
                ["tools/check_langmuir_build_env.py",
                 "Pre-build import-sanity check; mirrors "
                 "REQUIRED_LOCAL."],
            ])

    _heading(doc, "D.4 Build + test + installer flow", level=2)
    _para(doc, "EN — Developer commands:")
    _code(doc,
           "pip install -r requirements.txt\n"
           "pytest\n"
           "build.bat        # runs pyinstaller + optional inno-setup")
    _para(doc,
           "The full workflow is: (1) tools/check_langmuir_build_env.py "
           "validates every required local module can be imported; "
           "(2) pytest runs the ~1000-test suite; (3) "
           "python -m PyInstaller LangmuirMeasure.spec freezes the "
           "app into dist/LangmuirMeasure/; (4) ISCC.exe compiles "
           "LangmuirMeasure_setup.iss into installer_output/"
           "LangmuirMeasure_v3.0_setup.exe; (5) the installer "
           "warns at runtime if no system VISA is present and "
           "optionally chains vc_redist.x64.exe.")

    _heading(doc, "D.5 Runtime prerequisites", level=2)
    _para(doc,
           "Python is not required on the target PC — PyInstaller "
           "freezes it in.  The operator must install a system "
           "VISA (Keysight IO Libraries or NI-VISA) once; that "
           "also provides the GPIB driver for the USB adapter.  "
           "See docs/INSTALL_prereqs.md for the full checklist and "
           "the operator install sequence.")

    _heading(doc, "D.6 Project conventions", level=2)
    _bullet(doc,
             "Pure-Python analysis modules NEVER import Qt at "
             "module load.")
    _bullet(doc,
             "Each analysis result dict carries explicit "
             "fit_status / fit_error_reason / fit_warning_reason "
             "fields (see dlp_fit_models.FitStatus).")
    _bullet(doc,
             "VISA exceptions are raised from drivers as "
             "ClassifiedVisaError and surfaced via "
             "visa_errors.format_for_operator() in the GUI.")
    _bullet(doc,
             "CSV writers prepend the schema banner from "
             "dlp_csv_schema.write_header().")
    _bullet(doc,
             "Analyses write a JSON sidecar next to the CSV via "
             "analysis_options_sidecar.write_sidecar().")
    _bullet(doc,
             "Tests use offscreen Qt "
             "(QT_QPA_PLATFORM=offscreen) and a per-test "
             "deleteLater drain to avoid Qt teardown crashes.")

    _heading(doc, "D.7 Where to start for future development",
              level=2)
    _bullet(doc,
             "Need to change science? Touch only the pure modules "
             "(dlp_single_analysis, dlp_double_analysis, "
             "dlp_fit_models) and extend the status schema.")
    _bullet(doc,
             "Need a new instrument? Mirror the pattern in "
             "keysight_b2901.py + fake_b2901.py; raise "
             "ClassifiedVisaError from connect().")
    _bullet(doc,
             "Need a new UI option? Extend the relevant "
             "*Options dataclass, add to the dialog, persist "
             "through to_dict/from_dict.")
    _bullet(doc,
             "Need a new result field? Extend the result dict, "
             "the HTML renderer, the plain-text history line, and "
             "the sidecar summary in one coordinated edit.")
    _bullet(doc,
             "Want to run the frozen EXE from CI? See "
             "LangmuirMeasure.spec + build.bat; one open task is "
             "a frozen-binary smoke test.")

    _heading(doc, "D.8 Known limitations / future work", level=2)
    _bullet(doc,
             "n_i CI is currently fit-only; adding a full error "
             "budget for probe area and ion mass is a future pass.")
    _bullet(doc,
             "A Triple-probe Help dialog and an Options dialog "
             "analogous to Single / Double are not yet shipped.")
    _bullet(doc,
             "SCPI_ERROR classification is reserved in the "
             "VisaErrorKind enum but not yet populated (needs "
             "SYST:ERR? polling).")
    _bullet(doc,
             "Two pre-existing splitter-pixel-ratio tests are "
             "fragile on high-DPI monitors; a pre-existing "
             "Qt-offscreen teardown issue also affects the "
             "instrument-options dialog tests.")
    _bullet(doc,
             "A heuristic clipping detector exists for legacy "
             "CSVs; it is deliberately conservative (near-bit-"
             "identical plateau) and may miss noisy clamped "
             "data.  Output is always labelled 'suspected'.")

    _page_break(doc)

    # =======================================================================
    # Part E — Glossary / Glossar
    # =======================================================================
    _heading(doc, "E. Glossary / Glossar", level=1)

    gl = [
        ("Langmuir probe / Langmuir-Sonde",
         "A small metal electrode inserted into a plasma.  The "
         "current drawn at a swept voltage reveals plasma "
         "parameters.",
         "Eine kleine Metall-Elektrode, die ins Plasma ragt.  "
         "Der Strom bei variierter Spannung liefert Plasma-"
         "Parameter."),
        ("Single probe / Einfachsonde",
         "One tip plus a chamber-wall reference.  Measures "
         "V_f, T_e, V_p, n_e.",
         "Eine Spitze plus Kammerreferenz.  Misst V_f, T_e, "
         "V_p, n_e."),
        ("Double probe / Doppelsonde",
         "Two identical floating tips; no chamber reference "
         "needed.  Measures T_e and I_sat.",
         "Zwei baugleiche, schwebende Spitzen; keine "
         "Kammerreferenz nötig.  Misst T_e und I_sat."),
        ("Triple probe / Dreifachsonde",
         "Three tips, one actively biased.  Derives T_e and "
         "n_e from the ratio of two potential differences "
         "at each tick.",
         "Drei Spitzen, eine aktiv vorgespannt.  T_e und n_e "
         "aus dem Verhältnis der Potentialdifferenzen pro "
         "Zeitschritt."),
        ("Floating potential V_f",
         "The probe voltage at which the net current is zero.  "
         "Depends on the balance of electron and ion flux.",
         "Die Sondenspannung bei der der Nettostrom Null ist. "
         "Ergibt sich aus der Balance von Elektronen- und "
         "Ionenfluss."),
        ("Plasma potential V_p",
         "The potential of the bulk plasma.  Between V_f and "
         "V_p the current is dominated by retarded electrons.",
         "Das Potential des Plasmas.  Zwischen V_f und V_p "
         "überwiegen die retardierten Elektronen."),
        ("Electron temperature T_e",
         "Kinetic temperature of the electrons, in eV.  Extracted "
         "from the slope of the semilog I–V in the retarding zone.",
         "Kinetische Temperatur der Elektronen in eV.  Aus der "
         "Steigung des semilog-I–V in der Retardierungszone."),
        ("Ion saturation current I_sat",
         "The current plateau on the ion side of a floating probe, "
         "proportional to n_i · A · v_Bohm.",
         "Der Strom-Plateauwert auf der Ionenseite, proportional "
         "zu n_i · A · v_Bohm."),
        ("Electron density n_e / Ion density n_i",
         "Number density of the species, typically in m⁻³.  "
         "Derived from I_sat, probe area, T_e, and ion mass via "
         "the Bohm flux.",
         "Teilchendichte in m⁻³.  Aus I_sat, Sondenfläche, T_e "
         "und Ionenmasse über den Bohm-Fluss."),
        ("Confidence interval (95 %) / Konfidenzintervall",
         "A range that would contain the true value 95 out of "
         "100 times under repeated measurement.  Narrower is "
         "better only if the method is honest — see fit-only "
         "caveat.",
         "Bereich, der bei Wiederholung in 95 von 100 Fällen den "
         "wahren Wert enthält.  Schmaler ist nur dann besser, "
         "wenn die Methode ehrlich ist — siehe „fit-only“."),
        ("Compliance / Clipping",
         "When the SMU hits its programmed current limit.  The "
         "instrument clamps I to a constant value, which looks "
         "like a plateau in the I–V curve and biases the fit.",
         "Wenn die SMU ihr programmiertes Stromlimit erreicht. "
         "Das Gerät klemmt I auf einen konstanten Wert; in der "
         "Kennlinie sieht das wie ein Plateau aus und verfälscht "
         "den Fit."),
        ("Bootstrap CI",
         "A non-parametric method: resample the residuals of the "
         "base fit many times, refit each time, and take "
         "percentiles of the resulting T_e values.",
         "Nicht-parametrisches Verfahren: Residuen des Basis-Fits "
         "mehrfach neu ziehen, jedes Mal erneut fitten, "
         "Perzentile der T_e-Werte bilden."),
        ("VISA",
         "An industry-standard abstraction (visa32.dll / "
         "visa64.dll) that lets software talk to instruments over "
         "GPIB / USB / LAN / RS232 with one API.",
         "Industriestandard-Abstraktion (visa32.dll / "
         "visa64.dll), die einheitlichen Zugriff auf Instrumente "
         "über GPIB / USB / LAN / RS232 erlaubt."),
        ("GPIB",
         "IEEE-488 parallel bus.  Reliable, addressable, up to "
         "15 instruments per bus.  Needs a hardware adapter "
         "(Keysight 82357B, NI GPIB-USB-HS).",
         "IEEE-488 Parallel-Bus.  Zuverlässig, adressierbar, "
         "bis 15 Geräte pro Bus.  Braucht einen Hardware-Adapter "
         "(Keysight 82357B, NI GPIB-USB-HS)."),
        ("RS232",
         "Legacy serial link (COM port).  Slower than GPIB but "
         "needs no special adapter on a native serial port.  "
         "Used by the K2000 as a fallback.",
         "Klassischer serieller Bus (COM-Port).  Langsamer als "
         "GPIB, aber ohne Spezialadapter nutzbar.  Vom K2000 "
         "als Fallback verwendet."),
        ("Sidecar file / Sidecar-Datei",
         "A small JSON file written next to a measurement CSV "
         "that records the exact analysis options used.  Enables "
         "reproducible re-analysis.",
         "Kleine JSON-Datei neben einer Mess-CSV, die die "
         "verwendeten Analyse-Optionen festhält.  Ermöglicht "
         "reproduzierbare Re-Analysen."),
        ("Fit status",
         "A small taxonomy (OK, POOR, WARNING, "
         "non_converged, bounds_error, insufficient_data, "
         "bad_input, numerical_error) that tells the operator "
         "whether the numbers on screen can be trusted.",
         "Kleine Taxonomie (OK, POOR, WARNING, non_converged, "
         "bounds_error, insufficient_data, bad_input, "
         "numerical_error), die anzeigt, ob die angezeigten "
         "Zahlen vertrauenswürdig sind."),
        ("NPLC",
         "Number of power-line cycles used as integration time "
         "for a digitiser reading.  Lower = faster + noisier; "
         "higher = slower + quieter.  1.0 is ≈ 20 ms at 50 Hz.",
         "Number of Power Line Cycles — Integrationszeit als "
         "Vielfaches der Netzperiode.  Niedriger = schneller + "
         "rauschiger; höher = langsamer + ruhiger.  1,0 ≈ 20 ms "
         "bei 50 Hz."),
        ("Bohm velocity v_Bohm",
         "The ion sound speed √(kT_e/m_i); governs how fast ions "
         "enter the sheath around a probe and sets the link from "
         "I_sat to n_i.",
         "Die Ionen-Schallgeschwindigkeit √(kT_e/m_i); bestimmt "
         "die Ionen-Eintrittsrate in die Randschicht und damit "
         "den Zusammenhang zwischen I_sat und n_i."),
        ("Huber loss / Huber-Verlust",
         "A robust fit-loss that behaves like squared error near "
         "zero and like absolute error far from zero.  Resistant "
         "to occasional outliers (e.g. clipped points).",
         "Robuste Fit-Verlustfunktion, die nahe Null wie "
         "quadratischer und weit entfernt wie absoluter Fehler "
         "wirkt.  Widerstandsfähig gegen einzelne Ausreißer "
         "(z. B. geclippte Punkte)."),
    ]
    _table(doc,
            ["Term / Begriff",
             "English explanation",
             "Deutsche Erklärung"],
            [[term, en, de] for term, en, de in gl])

    _page_break(doc)

    # ---- back matter ----
    _heading(doc, "F. Contact / Kontakt", level=1)
    _para(doc,
           "EN — Issues, feature requests, and patches: open a "
           "ticket in the project's GitHub issue tracker.  For "
           "bench-operator questions at JLU-IPI, contact the "
           "I. Physikalisches Institut.")
    _para(doc,
           "DE — Fehlerberichte, Feature-Wünsche und Patches: "
           "Ticket im GitHub-Issue-Tracker des Projekts anlegen. "
           "Für Fragen zum Laborbetrieb am JLU-IPI bitte an das "
           "I. Physikalische Institut wenden.")

    out = Path(__file__).resolve().parent \
          / "LangmuirMeasure_Documentation.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    size_kb = os.path.getsize(path) / 1024.0
    print(f"Wrote {path} ({size_kb:.1f} KB)")
